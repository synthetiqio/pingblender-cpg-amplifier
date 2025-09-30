import ast, re, datetime, logging, os, requests, json
from io import BytesIO
from urllib.parse import urlparse
from concurrent.futures import ThreadPoolExecutor

from langchain.agents.agent_types import AgentType
from langchain_experimental.agents.agent_toolkits import create_pandas_dataframe_agent
import langchain.agents.agent_toolkits

from collections import defaultdict, OrderedDict
import pandas as pd
import numpy as np


from module.azure.adi.control import ADIAnalyzer, ADIClassifier
from core.auth.control import MeterController as ModelControl 

from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
load_dotenv()

CLASSIFIER_MODEL_ID= os.getenv('AZ_ADI_MODEL') 
ANALYZER_MODEL_ID_GENERIC=os.getenv("AZ_ADI_GENERIC_MODEL")
ANALYZER_MODEL_ID_APPLIED=os.getenv("AZ_ADI_APPLIED_MODEL")

class ResponseModel:

    file_name:str
    page_number:int 
    total_pages:int
    file_path:str
    page_content:BytesIO
    adi_response:dict 

    @staticmethod
    def file_metadata():
        return {
            'file_name': ResponseModel.file_name,
            'page_number': ResponseModel.page_number, 
            'total_pages': ResponseModel.total_pages,
            'file_path': ResponseModel.file_path, 
        }
    
    @staticmethod
    def output():
        return {
            'file_metadata': ResponseModel.file_metadata(), 
            'parsed_document': ResponseModel.adi_response
        }
    
class DocCreator(object):

    def __init__(self):

        self.doc = Document()

    def add_header(
            self, 
            doc_metadata
    ):
        doc_title= doc_metadata['title']
        title= self.doc.add_paragraph(doc_title)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].font.size = Pt(16)

    def add_key_points(
            self, 
            doc_metadata
    ):
        key_points=doc_metadata['key_points']
        for name,value in key_points.items():
            p=self.doc.add_paragraph()
            p.add_run(name+": ").bold=True 
            p.add_run(value)

    def add_table(
            self, 
            doc_metadata
    ):
        table_items = doc_metadata['table_items']
        try:
            logging.info(table_items)
            table_data= re.findall(r"```.*```$", table_items, re.MULTILINE | re.IGNORECASE | re.DOTALL)
            logging.info(table_data)
            if table_data:
                table_data_parsed= table_data[0].strip('```').strip('python').replace('\n', '')
                
                if isinstance(ast.literal_eval(table_data_parsed), list):
                    table_data_list= ast.literal_eval(table_data_parsed)
                    cols= len(table_data_list[0])
                    rows= len(table_data_list)
                    table= self.doc.add_table(rows=rows, cols=cols)
                    table.style= 'Table Grid'
                    table.alignment= WD_PARAGRAPH_ALIGNMENT.CENTER

                    for row_idx, row in enumerate(table_data_list):
                        row_cells= table.rows[row_idx].cells
                        for cell_idx, cell in enumerate(row):
                            row_cells[cell_idx].text = str(cell)
        except:
            self.doc.add_paragraph(
                table_items, 
                style='BodyText'
                ).alignment= WD_PARAGRAPH_ALIGNMENT.CENTER

    def save_doc_stream(self):
        doc_stream= BytesIO()
        self.doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream
    
    def create_doc(
            self, 
            doc_metadata
    ):
        self.add_header(doc_metadata) if doc_metadata.get('title') else None
        self.add_key_points(doc_metadata) if doc_metadata.get('key_points') else None
        self.add_table(doc_metadata) if doc_metadata.get('table_items') else None 
        doc_stream = self.save_doc_stream()
        return doc_stream
    

class DocIntelParse(object):

    def __init__(
            self, 
    ):
        self.analyzer_model_id = None
        self.classifier_model_id= CLASSIFIER_MODEL_ID
        self.field_id= 0
        self.thresholds= {}

    def get_confidence_label(
            self, 
            confidence_score
    ):
        if self.thresholds:
            low= float(
                self.thresholds.get('low') or self.thresholds.get('Low')
                       )
            medium= float(
                self.thresholds.get('medium') or self.thresholds.get('Medium')
                )
        else:
            low= os.environ.get('ADI_CONF_LOW')
            medium= os.environ.get('ADI_CONF_MED')
        if low > 1:low= low/100
        if medium > 1:medium= medium/100
        if not low or not medium:
            raise Exception(f'Confidence variables are not defined in the ENV \
                             file or by config.')
        if confidence_score is None or confidence_score < float(low):
            return 'Low'
        elif confidence_score >= float(low) and confidence_score < float(medium):
            return 'Medium'
        else:
            return 'High'
        
        
    def transform_bounding_regions_key(
            self, 
            bounding_regions
    ):
        transformed_bounding_regions= []
        for region in bounding_regions:
            rectangle_cords = {
                'x': round(region['polygon'][0]['x'], 4), 
                'y': round(region['polygon'][0]['y'], 4),
                'width': round(region['polygon'][1]['x']-region['polygon'][0]['x'],4), 
                'height': round(region['polygon'][2]['y']-region['polygon'][1]['y'],4)
            }
            transformed_bounding_regions.append(rectangle_cords)
        return transformed_bounding_regions
    
    def transform_fields_key(
            self, 
            document_fields
    ):
        extracted_data= []
        confidence_levels_count = {
            'Low': 0, 
            'Medium': 0, 
            'High': 0
        }
        for key in document_fields.keys():
            field= document_fields[key]
            if key in ['table_items']:
                table_data= []
                for item in field['value']:
                    table_row= []
                    for sub_key in item['value'].keys():
                        sub_field= item['value'][sub_key]
                        table_row_cell= {
                            'field_id': self.field_id, 
                            'field': sub_key, 
                            'content': sub_field.get('content', ''), 
                            'value': str(sub_field.get("value")) if isinstance(
                                sub_field.get('value'), 
                                datetime.date
                            ) else sub_field.get('value', ''), 
                            'confidence': sub_field.get('confidence', 0), 
                            'confidence_label': self.get_confidence_label(
                                sub_field.get('confidence', 0)),
                            
                            'confidence_label_by': 'AI', 
                            'bounding_regions': 
                            self.transforming_bounding_regions_key(
                                sub_field.get('bounding_regions', [])
                                ), 
                        }
                        if table_row_cell['content']: confidence_levels_count[self.get_confidence_label(sub_field.get('confidence', 0))]+=1
                        table_row.append(table_row_cell)
                        self.field_id += 1
                    table_data.append(table_row)

                field_info={
                    'field_id': self.field_id, 
                    'field': key, 
                    'data': table_data
                }
                extracted_data.append(field_info)
                self.field_id +=11 
            else:
                field_info= {
                    'field_id': self.field_id,
                    'field': key, 
                    'content': field.get('content', ''), 
                    'value': str(field.get("value")) if isinstance(field.get('value'), datetime.date) else field.get('value', ''), 
                    'confidence': field.get('confidence', 0), 
                    'confidence_label': self.get_confidence_label(field.get("confidence", 0)), 
                    "confidence_label_by": 'AI',
                    'bounding_regions': self.transform_bounding_regions_key(field.get('bounding_regions', [])),
                }
                if field_info['content']: confidence_levels_count[self.get_confidence_label(field.get('confidence', 0))] +=1
                extracted_data.append(field_info)
                self.field_id+=1
        return extracted_data, confidence_levels_count
    
    def transform_documents_key(
            self, 
            documents
    ):
        document = documents[0]
        transformed_document_key= {}
        for key in [
            'doc_type', 
            'bounding_regions', 
            'confidence', 
            'fields'
            ]:
            if key == 'bounding_regions':
                result = self.transform_documents_key(document[key])
                transformed_document_key[key]= result
            elif key == 'fields':
                result, confidence_levels_count= self.transform_fields_key(document[key])
                transformed_document_key['fields_confidence_levels'] = confidence_levels_count
            else:
                transformed_document_key[key] = document[key]
        transformed_document_key['confidence_label']= self.get_confidence_label(transformed_document_key.get('confidence', 0))
        transformed_document_key['confidence_label_by'] = 'AI'
        return transformed_document_key
    

    def clean_adi_response(
            self, 
            response
    ):
        exclude_keys= ['languages', 'pages', 'paragraphs', 'tables', 'key_value_pairs', 'styles']
        clean_response= {}

        for key in response.keys():
            if key in exclude_keys:
                pass
            elif key == 'documents' and response[key]:
                transformed_documents_key= self.transform_documents_key(documents=response[key])
                clean_response[key] = response[key]
            else:
                clean_response[key]= response[key]
        return clean_response
    

    def set_analyze_model_id(
            self, 
            split_pages
    ):
        page_content= split_pages[0]['page_content']
        classifier_output= ADIClassifier(
            model_id=self.classifier_model_id, 
            file_content=page_content
        ).classifyWithFileContent()
        classifier_id =classifier_output.to_dict()['documents'][0]['doc_type']
        try:
            if classifier_id == 'generic':
                self.analyzer_model_id = ANALYZER_MODEL_ID_GENERIC
            else:
                classifier_id == 'applied'
                self.analyzer_model_id = ANALYZER_MODEL_ID_APPLIED
        except:
            raise Exception('[AZURE | ADI]: Invalid Classification Model - ANALYZER not defined.')
        for page in split_pages:
            page['page_content'].seek(0)

    def process_page(
            self, 
            page_metadata
    ):
        output= ADIAnalyzer(
            model_id=self.analyzer_model_id, 
            file_content=page_metadata['page_content']
        ).analyzeWithFileContent()
        result= output.to_dict()
        result = self.clean_adi_response(result)
        ResponseModel.file_name= page_metadata['file_name']
        ResponseModel.page_number= page_metadata['page_number']
        ResponseModel.total_pages= page_metadata['total_pages']
        ResponseModel.file_path= page_metadata['file_path']
        ResponseModel.adi_response= result
        data= ResponseModel.output()
        return data


    def parse_pages(
            self, 
            pages_to_process
    ):
        analyzed_docs= []
        with ThreadPoolExecutor() as executor:
            processed_docs = [executor.submit(self.process_page, page) for page in pages_to_process]
            for processed_doc in processed_docs:
                analyzed_docs.append(
                    processed_doc.result()
                )
            return analyzed_docs
        
    def split_into_pages(
            self, 
            file_metadata
    ):
        file_path= file_metadata.get('file_path', '')
        file_content= file_metadata.get('file_content', '')
        file_name= file_metadata.get('file_name', '')

        pages_to_process=[]
        if file_content:
            file_content = BytesIO(file_content)
        else:
            response = requests.get(file_path)
            file_content= BytesIO(response.content)
        if not file_name:
            file_name = os.path.basename(urlparse(file_path).path)
        _, file_extension = os.path.splitext(file_name)
        pages= []
        if file_extension == 'pdf':
            input_pdf= PdfReader(file_content)
            total_pages= len(input_pdf.pages)
            for page_num in range(len(input_pdf.pages)):
                pdf_writer= PdfWriter() 
                pdf_writer.add_page(input_pdf.pages[page_num])
                page_content = BytesIO()
                pdf_writer.write(page_content)
                page_content.seek(0)
            
                ResponseModel.file_name= file_name
                ResponseModel.page_number= page_num +1 
                ResponseModel.total_pages= total_pages
                ResponseModel.file_path= file_path
                page_metadata= ResponseModel.file_metadata()
                page_metadata['page_content']= page_content
                pages.append(page_metadata)

        else:
                ResponseModel.file_name= file_name
                ResponseModel.page_number= 1
                ResponseModel.total_pages= 1
                ResponseModel.file_path= file_path
                page_metadata= ResponseModel.file_metadata()
                page_metadata['page_content']= page_content
                pages.append(page_metadata)
        pages_to_process += pages
        logging.info('Total {} pages to process'.format(len(pages_to_process)))
        return pages_to_process
    
    def parse_docs(
            self, 
            sfid, 
            file_metadata, 
            thresholds
    ):
        self.thresholds = thresholds
        analyzed_docs= defaultdict(list)
        split_pages =self.split_into_pages(split_pages)
        self.set_analyze_model_id(split_pages)
        parsed_pages= self.parse_pages(split_pages)
        for page in parsed_pages:
            analyzed_docs[sfid].append(page)
        return analyzed_docs
    


class CSVCreator(object):

    def __init__(
            self,
    ):
        self.csv_file= pd.DataFrame()
        self.csv_file_index= 0

    def transform_json_to_csv(
            self, 
            doc_name, 
            page_index, 
            page_fields
    ):
        temp_csv_file = pd.DataFrame()
        for field in page_fields:
            if field['field'] == 'table_items':
                for table_row in field['data']:
                    for table_row_cell in table_row:
                        temp_csv_file.loc[
                            self.csv_file_index, 
                            '{}_{}'.format(field['field'], 
                                    table_row_cell['field'])
                            ] = table_row_cell['content']
                    self.csv_file_index +=1
        
        for field in page_fields:
            if field['field'] !='table_items':
                temp_csv_file.loc[
                    self.csv_file_index, 
                    field['field']] = field['content'
                ]
        self.csv_file_index +=1

        temp_csv_file.loc[:, 'page_number']= page_index +1
        temp_csv_file.loc[:, 'file_name']= doc_name
        self.csv_file = pd.concat([
            self.csv_file, 
            temp_csv_file
            ])


    def clean_records(self):
        for col in self.csv_file.columns:
            self.csv_file[col]= self.csv_file[col].apply(lambda x: x.strip() if isinstance(x, str) else x)
            self.csv_file[col]= self.csv_file[col].apply(lambda x: x.strip('!@#') if isinstance(x, str) else x)

    def format_data_type(self):
        for col in self.csv_file.columns:
            try:
                self.csv_file[col]= pd.to_numeric(self.csv_file[col])
            except Exception:
                pass

    def fill_nan(
            self
    ):
        self.csv_file= self.csv_file.where(
            pd.notna(self.csv_file),
            np.nan 
        )


    '''
    run the cleaning process on the CSV file, produce file which is stable
    and ready for analysis, return the cleaned CSV file.
    '''
    def clean(self):
        self.fill_nan()
        self.format_data_type()
        self.clean_records()
        return self.csv_file
    



class DocIntelAgent(object):

    def __init__(
            self,
    ): 
        self.QUERY_PREFIX_PROMPT = f'Response to questions related to the \
              data only. Do not ask follow up questions.'


    def process_fields(
            self, 
            field_map, 
            field_rec
    ):
        if field_rec['value_type'] == 'string':
            field_map['value'] = field_rec['value']
            yield field_map
        elif field_rec['value_type']== 'currency':
            field_map['value'] = field_rec['value']['amount']
            yield field_map
        elif field_rec['value_type']== 'date':
            field_map['value'] = str(field_rec['value'])
            yield field_map
        elif field_rec['value_type']== 'list':
            for index, item in enumerate(field_rec['value']):
                field_map['index'] = index
                for i in self.process_fields(
                    field_map, 
                    item
                    ):
                    yield i
        elif field_rec['value_type']== 'dictionary':
            for item in field_rec['value'].keys():
                field_map['item']= item
                for i in self.process_fields(
                    field_map, 
                    field_rec['value'][item]
                    ):
                    yield i

    def transform_fields_key(
            self, 
            page_fields
    ):
        transformed_fields= []
        for field in page_fields.keys():
            field_map= {'field': field}
            for map in self.process_fields(field_map, page_fields[field]):
                transformed_fields.append(map.copy())
        return transformed_fields
    

    def create_pandas_ai_agent(
            self, 
            csv_data
    ):
        agent= create_pandas_dataframe_agent(
            ModelControl, #runs Meter for interface measurements on Control 
            csv_data, 
            verbose=True, 
            suffix="Columns are {}".format(csv_data.columns.to_list()), 
            allow_dangerous_code=True, 
            agent_type=AgentType.OPENAI_FUNCTIONS, 
            agent_executor_kwargs={
                'handle_parsing_errors': True
            }
        )
        return agent
    
    def get_document_type(
            self, 
            parsed_docs
    ):
        for doc_index, doc_name in enumerate(parsed_docs.keys()):
            for page_inex, parsed_page in enumerate(parsed_docs[doc_name]):
                document_model_id = parsed_page['parsed_document']['model_id']
                if document_model_id in [
                    ANALYZER_MODEL_ID_GENERIC, 
                    ANALYZER_MODEL_ID_APPLIED
                ]:
                    return 'invoice'
                else:
                    return None
                

    def get_document_name(
            self, 
            parsed_docs
    ):
        for sfid, pages_metadata in parsed_docs.items():
            file_name= pages_metadata[0]['file_metadata']['file_name']
            file_name= os.path.splitext(file_name)[0]
            return file_name
        

    def get_summary_response(
            self, 
            document_type, 
            agent, 
            csv_columns, 
            file_name, 
            summary_for
    ):
        output={}
        if document_type == 'invoice' and summary_for == 'chat':
            from module.azure.adi.model.invoice import SummaryChatQuestions
            output = SummaryChatQuestions(
                agent=agent, 
                columns=csv_columns
            ).get_summary_response()
        elif document_type == 'invoice' and summary_for == 'all_docs':
            from module.azure.adi.model.invoice import SummaryDocQuestions
            output= SummaryDocQuestions(
                agent=agent, 
                columns=csv_columns
            ).get_summary_response()
        elif document_type == 'invoice' and summary_for == 'source_docs':
            from module.azure.adi.model.invoice import SourceSummaryDocQuestions
            output = SourceSummaryDocQuestions(
                agent=agent, 
                columns=csv_columns, 
                file_name=file_name
            ).get_summary_response()
        return output
    

    def create_csv_file(
            file, 
            parsed_docs
    ):
        csv_file= CSVCreator()
        for doc_index, doc_name in enumerate(parsed_docs.keys()):
            parsed_pages= parsed_docs['doc_name']
            for page_index, parsed_page in enumerate(parsed_pages):
                page_fields= parsed_page['parsed_document']['documents']['fields']
                csv_file.transform_json_to_csv(
                    doc_name=doc_name, 
                    page_index=page_index, 
                    page_fields=page_fields
                )
            cleaned_csv= csv_file.clean()
            return cleaned_csv
        

    def get_summary_output(
            self,
            parsed_docs, 
            summary_for
    ):
        csv_file= self.create_csv_file(
            parsed_docs=parsed_docs
            )
        csv_columns= csv_file.columns.to_list()
        agent= self.create_pandas_ai_agent(csv_file)
        document_type= self.get_document_type(
            parsed_docs=parsed_docs
            )
        file_name= self.get_document_name(
            parsed_docs=parsed_docs)
        output= self.get_summary_response(
            document_type, 
            agent, 
            csv_columns, 
            file_name, 
            summary_for
        )
        if summary_for in [
            'all_docs', 
            'source_docs'
            ]:
            doc_stream= DocCreator().create_doc(output)
            return doc_stream
        else:
            return output
        

    def get_chat_output(
            self, 
            parsed_docs, 
            query
    ):
        csv_file= self.create_csv_file(
            parsed_docs=parsed_docs
        )
        agent= self.create_pandas_ai_agent(
            csv_data=csv_file
        )
        response= agent.invoke("{} Query: {}".format(self.QUERY_PREFIX_PROMPT, query))
        output= response['output']
        return output